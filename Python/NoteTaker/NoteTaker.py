import pyPdf;
import random;

from docx import Document;

####################################################################
# Written by Ryan D'souza
# Generates notes randomly from a PDF
# See: github.com/dsouzarc/dotfiles/tree/master/Python/NoteTaker
####################################################################

#UPDATE THESE
fileName = "Dubliners.pdf";
noteSheetName = "Dubliners Notes.md";
noteSheetWordName = "Dubliners Notes.docx";
startPage = 4; #The actual first page of the book

#Optional Update
noteFrequency = 2;
minWordsInLine = 5;

failCounter = 0;

#Returns a random note line based on some conditions
def randomNote(lines):

    global failCounter;

    #If, within 6 tries, we cannot find a sentence that matches what we want, just return
    if failCounter > 5:
        failCounter = 0;
        return None;

    #Choose a random sentence
    randomNoteLine = random.choice(lines);

    #If it doesn't match the conditions we want, recurse until we find a sentence that does
    if len(randomNoteLine.split(" ")) <= minWordsInLine:
        failCounter += 1;
        return randomNote(lines);

    #Otherwise, just return the line we found
    else:
        failCounter = 0;
        return randomNoteLine + ".";

#Markdown format line for notes
def formatLine(pageNumber, line):
    formattedLine = "- **Page " + str(pageNumber) + "**: " + line + "\n";
    return formattedLine;

#PDF version of the book
book = file(fileName, "rb");
pdf = pyPdf.PdfFileReader(book);

if pdf.isEncrypted:
    pdf.decrypt('');

#Create a notes file
notes = open(noteSheetName, "w");
wordDocument = Document();

for i in range(startPage, pdf.getNumPages()):

    #If it is time to take a note
    if (i - startPage) % noteFrequency == 0:

        #Gets each line on the page
        page = pdf.getPage(i);
        text = page.extractText();
        text = text.replace("\n", " ");
        lines = text.split(".");

        failCounterLocal = 0;

        #While loop for continually retrying 'try' until there is no exception
        while True:

            #If we keep on getting errors (like encoding errors), try the prior page
            if failCounterLocal > 5:
                page = pdf.getPage(i - 1);
                text = page.extractText();
                text = text.replace("\n", " ");
                lines = text.split(".");
                failCounterLocal = 0;

            #Try writing to both files (possible encoding error)
            try: 

                #Get a formatted line at random from the page
                randomNoteLine = randomNote(lines);

                #If we can't find a line that satisfies our conditions, reset to the for-loop which will give us a new page and new lines
                if randomNoteLine is None:
                    break;

                #Markdown formatted line
                formattedLine = formatLine(i + 1, randomNoteLine);

                #Add the line to the word document note sheet and the markdown note sheet
                notes.write(formattedLine);
                noteLine = wordDocument.add_paragraph('', style='ListBullet');
                noteLine.add_run('Page ' + str(i + 1) + ': ').bold = True
                noteLine.add_run(randomNoteLine).bold = False;
                break;

            #Sometimes there is an encoding error. Increment our problem counter. Head back to while
            except (UnicodeError):
                failCounterLocal += 1;

#Save and print mission success :)
notes.close();
wordDocument.save(noteSheetWordName);
print("Notes created. Check: " + noteSheetName + " and " + noteSheetWordName);
